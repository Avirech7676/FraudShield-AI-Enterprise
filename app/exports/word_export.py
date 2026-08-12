from docx import Document


class WordExporter:

    @staticmethod
    def export(text, filename):

        doc = Document()

        doc.add_heading(

            "FraudShield AI Enterprise",

            level=1

        )

        doc.add_paragraph(text)

        doc.save(filename)
